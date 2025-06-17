import streamlit as st
from deep_translator import GoogleTranslator
from docx import Document
import tempfile
import os
import fitz  # PyMuPDF
from pdf2image import convert_from_bytes
import pytesseract
from PIL import Image


def translate_text(text, dest_lang='hi'):
    try:
        if len(text.strip()) == 0:
            return ""
        return GoogleTranslator(source='auto', target=dest_lang).translate(text)
    except:
        return text  # fallback if translation fails

def process_docx_preserve_format(file, lang_code):
    doc = Document(file)
    new_doc = Document()

    for para in doc.paragraphs:
        new_para = new_doc.add_paragraph()
        for run in para.runs:
            original_text = run.text
            translated_text = translate_text(original_text, lang_code)
            new_run = new_para.add_run(translated_text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
        new_para.style = para.style

    output_path = os.path.join(tempfile.gettempdir(), "translated_preserved.docx")
    new_doc.save(output_path)
    return output_path

def is_pdf_text_based(file_bytes):
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            if page.get_text().strip():
                return True
        return False
    except:
        return False

def process_pdf_translate(file, lang_code):
    file_bytes = file.read()
    is_text_pdf = is_pdf_text_based(file_bytes)
    new_doc = Document()

    if is_text_pdf:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text = page.get_text()
            for line in text.split("\n"):
                if line.strip():
                    translated_line = translate_text(line, lang_code)
                    new_doc.add_paragraph(translated_line)
    else:
        images = convert_from_bytes(file_bytes)
        for image in images:
            text = pytesseract.image_to_string(image, lang='eng')
            for line in text.split("\n"):
                if line.strip():
                    translated_line = translate_text(line, lang_code)
                    new_doc.add_paragraph(translated_line)

    output_path = os.path.join(tempfile.gettempdir(), "translated_pdf_output.docx")
    new_doc.save(output_path)
    return output_path

# Streamlit UI
st.title("📄 Document Translator with Format Preservation")
uploaded_file = st.file_uploader("Upload a .docx or .pdf file", type=["docx", "pdf"])
lang = st.selectbox("Select Output Language", options=[
    ("English", "en"),
    ("Hindi", "hi"),
    ("Marathi", "mr"),
    ("Gujarati", "gu"),
    ("Tamil", "ta"),
    ("Telugu", "te"),
    ("Bengali", "bn"),
    ("French", "fr"),
    ("German", "de"),
    ("Spanish", "es")
], format_func=lambda x: x[0])
lang_code = lang[1]

if uploaded_file:
    st.info(f"Processing file: {uploaded_file.name}")

    if uploaded_file.name.endswith(".docx"):
        output_path = process_docx_preserve_format(uploaded_file, lang_code)
    elif uploaded_file.name.endswith(".pdf"):
        output_path = process_pdf_translate(uploaded_file, lang_code)
    else:
        st.error("Unsupported file format")
        st.stop()

    st.success("Translation Complete")

    with st.expander("View Sample Input and Output"):
        st.subheader("🔹 Original Text Sample:")
        if uploaded_file.name.endswith(".docx"):
            original_doc = Document(uploaded_file)
            original_sample = "\n".join(p.text for p in original_doc.paragraphs[:3] if p.text.strip())
        else:
            uploaded_file.seek(0)
            file_bytes = uploaded_file.read()
            if is_pdf_text_based(file_bytes):
                pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
                original_sample = pdf_doc[0].get_text()
            else:
                images = convert_from_bytes(file_bytes)
                original_sample = pytesseract.image_to_string(images[0], lang='eng')
        st.text_area("Original Sample", original_sample.strip(), height=150)

        st.subheader("🔸 Translated Text Sample:")
        translated_doc = Document(output_path)
        translated_sample = "\n".join(p.text for p in translated_doc.paragraphs[:3] if p.text.strip())
        st.text_area("Translated Sample", translated_sample, height=150)

    st.download_button("Download Translated DOCX", data=open(output_path, "rb"), file_name="translated_output.docx")
